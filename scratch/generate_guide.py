import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def generate_docx():
    doc = Document()
    
    # Page setup
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    
    # Custom styles & colors
    COLOR_PRIMARY = RGBColor(16, 44, 87)    # Deep Navy
    COLOR_SECONDARY = RGBColor(100, 116, 139) # Slate Grey
    COLOR_DARK = RGBColor(30, 41, 59)       # Dark Slate
    COLOR_ACCENT = RGBColor(225, 112, 85)   # Amber/Orange
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title.add_run("BANGLADESH POWER GRID EXPLORER")
    run_title.font.name = "Georgia"
    run_title.font.size = Pt(24)
    run_title.font.bold = True
    run_title.font.color.rgb = COLOR_PRIMARY
    
    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = subtitle.add_run("The Layperson's Handbook to Bangladesh's Energy Metrics, Units, and Data Dashboard")
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(12)
    run_sub.font.italic = True
    run_sub.font.color.rgb = COLOR_SECONDARY
    
    # Horizontal line replacement
    p_sep = doc.add_paragraph()
    p_sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sep.add_run("—" * 50).font.color.rgb = COLOR_SECONDARY
    
    # SECTION 1: Welcome & Introduction
    h1 = doc.add_heading(level=1)
    run_h1 = h1.add_run("1. Welcome to the Power Grid Explorer")
    run_h1.font.name = "Georgia"
    run_h1.font.size = Pt(16)
    run_h1.font.bold = True
    run_h1.font.color.rgb = COLOR_PRIMARY
    
    p = doc.add_paragraph()
    p.add_run(
        "Energy is the lifeblood of Bangladesh's fast-growing economy. Every light bulb, factory machine, "
        "irrigation pump, and computer depends on a complex network of power plants, gas fields, LNG terminals, "
        "and transmission lines. To a layperson, however, the data behind this infrastructure looks like a soup "
        "of technical jargon, confusing abbreviations, and massive numbers.\n\n"
        "The "
    )
    p.add_run("Power Grid Explorer").bold = True
    p.add_run(
        " was built to bring transparency to this vital sector. It aggregates and visualizes messy public databases—such "
        "as daily reports from the Power Grid Company of Bangladesh (PGCB) and Petrobangla—into a single, clean, "
        "interactive portal. This guide is designed to be read before diving into the live tool, ensuring that "
        "anyone, from students to utility supervisors, can interpret grid statistics with ease."
    )
    
    # SECTION 2: Deciphering the Units & Acronyms
    h2 = doc.add_heading(level=1)
    run_h2 = h2.add_run("2. Deciphering the Units & Acronyms")
    run_h2.font.name = "Georgia"
    run_h2.font.size = Pt(16)
    run_h2.font.bold = True
    run_h2.font.color.rgb = COLOR_PRIMARY
    
    p2 = doc.add_paragraph()
    p2.add_run(
        "To understand the Grid Explorer, you must first understand the units used to measure electricity and fuel. "
        "Here are the primary terms explained simply:"
    )
    
    # Bullet points for units
    bp1 = doc.add_paragraph(style='List Bullet')
    r = bp1.add_run("MW (Megawatt): ")
    r.bold = True
    r.font.color.rgb = COLOR_DARK
    bp1.add_run(
        "This measures power (instantaneous capacity or rate of energy generation/demand). Think of it like the "
        "speedometer of a car, showing how fast electricity is being produced or consumed at a single moment. "
        "One Megawatt (MW) can power approximately 800 to 1,000 average homes in Bangladesh. For reference, "
        "Bangladesh's total peak power demand typically ranges between 13,000 MW and 16,500 MW."
    )
    
    bp2 = doc.add_paragraph(style='List Bullet')
    r = bp2.add_run("MKWh (Million Kilowatt-Hours): ")
    r.bold = True
    r.font.color.rgb = COLOR_DARK
    bp2.add_run(
        "This measures energy (the total volume of electricity generated or consumed over a period of time). "
        "Think of it like the odometer of a car, showing the total distance traveled. One Kilowatt-hour (kWh) "
        "is the energy used by a 1,000-watt appliance (like a microwave) running for one hour. A Million "
        "Kilowatt-hours (MKWh) represents 1,000 Megawatt-hours (MWh). The grid explorer uses MKWh to show "
        "the total cumulative energy generated during an entire day (e.g., 320 MKWh)."
    )
    
    bp3 = doc.add_paragraph(style='List Bullet')
    r = bp3.add_run("MMCFD (Million Standard Cubic Feet per Day): ")
    r.bold = True
    r.font.color.rgb = COLOR_DARK
    bp3.add_run(
        "This measures gas volume flow rate. Bangladesh relies heavily on natural gas to fuel its power plants. "
        "This unit tells us how much gas is flowing out of the ground or LNG terminals in a 24-hour day. "
        "For example, a supply of 2,600 MMCFD means 2.6 billion cubic feet of gas flowed into the national grid that day."
    )
    
    bp4 = doc.add_paragraph(style='List Bullet')
    r = bp4.add_run("BBL (Barrel of Oil): ")
    r.bold = True
    r.font.color.rgb = COLOR_DARK
    bp4.add_run(
        "The standard unit of volume for crude oil and liquid petroleum products. One BBL is exactly 42 US gallons "
        "or approximately 159 liters. It is used to track liquid fuel reserves, imports, and consumption for "
        "diesel and furnace oil-based power plants."
    )
    
    bp5 = doc.add_paragraph(style='List Bullet')
    r = bp5.add_run("Lakh and Crore: ")
    r.bold = True
    r.font.color.rgb = COLOR_DARK
    bp5.add_run(
        "Unlike Western systems that count in millions and billions, Bangladesh uses the subcontinental system. "
        "One Lakh (1,00,000) equals one hundred thousand. One Crore (1,00,00,000) equals ten million. The Grid Explorer "
        "formats financial costs in BDT (Bangladeshi Taka) using Crore (e.g., 180 Crore Tk. daily fuel cost) so "
        "that local policy analysts and citizens can easily relate to the figures."
    )
    
    # SECTION 3: Key Grid Metrics & What They Mean
    h3 = doc.add_heading(level=1)
    run_h3 = h3.add_run("3. Key Grid Metrics & What They Mean")
    run_h3.font.name = "Georgia"
    run_h3.font.size = Pt(16)
    run_h3.font.bold = True
    run_h3.font.color.rgb = COLOR_PRIMARY
    
    p3 = doc.add_paragraph()
    p3.add_run(
        "When looking at the dashboard, you will see five main KPIs. Here is what they track and why they matter:"
    )
    
    # Table for metrics
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Metric'
    hdr_cells[1].text = 'Unit'
    hdr_cells[2].text = 'Layperson Explanation'
    
    # Style table headers
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = COLOR_PRIMARY
        
    metrics_data = [
        ("System Demand", "MW", "The maximum amount of power consumers requested simultaneously. Typically spikes in the evening (7 PM to 10 PM) due to household cooling and lighting."),
        ("Peak Generation", "MW", "The maximum power the grid actually managed to generate and supply at the moment of highest demand."),
        ("Load Shedding", "MW / %", "The deficit. If demand is 15,000 MW but plants can only supply 13,500 MW, the remaining 1,500 MW is cut off (load shedding) to prevent the entire grid from collapsing (a blackout)."),
        ("System Loss", "%", "Electricity lost as heat as it travels through transmission wires and transformers from power plants to your home. Lower is better; typically around 8-10% in Bangladesh."),
        ("Renewable Share", "%", "The percentage of power generated from clean, renewable sources (Solar, Wind, Hydro) relative to fossil fuels. Currently a small but growing fraction.")
    ]
    
    for metric, unit, desc in metrics_data:
        row_cells = table.add_row().cells
        row_cells[0].text = metric
        row_cells[1].text = unit
        row_cells[2].text = desc
        row_cells[0].paragraphs[0].runs[0].font.bold = True
        
    # Spacing after table
    doc.add_paragraph()
    
    # SECTION 4: Deep Dive - Energy vs. Power
    h4 = doc.add_heading(level=1)
    run_h4 = h4.add_run("4. Deep Dive: Energy vs. Power & The 8-10% Loss Rule")
    run_h4.font.name = "Georgia"
    run_h4.font.size = Pt(16)
    run_h4.font.bold = True
    run_h4.font.color.rgb = COLOR_PRIMARY
    
    p4 = doc.add_paragraph()
    p4.add_run(
        "A common point of confusion is the difference between Peak Power (MW) and Daily Energy Generation (MKWh). "
        "Let us use a simple, real-world comparison to clarify this:\n\n"
    )
    
    # Blockquote style
    p_bq = doc.add_paragraph()
    p_bq.paragraph_format.left_indent = Inches(0.5)
    r_bq = p_bq.add_run(
        "Imagine the power grid is a tap filling a bucket. "
        "The flow rate of water coming out of the tap is the POWER (MW). "
        "The total volume of water collected in the bucket at the end of the day is the ENERGY (MKWh).\n"
        "If you turn the tap on full blast for 1 hour (high MW) and then turn it off, you might collect "
        "the same amount of water (MKWh) as leaving the tap at a slow trickle (low MW) for 24 hours. "
        "The grid explorer tracks both: the maximum flow rate during peak hours (MW) and the total "
        "electricity bucket collected at the end of the day (MKWh)."
    )
    r_bq.font.italic = True
    r_bq.font.color.rgb = COLOR_SECONDARY
    
    p_loss = doc.add_paragraph()
    p_loss.add_run("The Station Auxiliary Consumption Rule (8-10%): ").bold = True
    p_loss.add_run(
        "When looking at the dashboard, you might notice that the sum of fuel generation (Gross Generation) is "
        "larger than the Net Grid Dispatch. Why? Because power plants require electricity to run their own heavy "
        "machinery, water pumps, cooling towers, and control rooms. This internal consumption typically eats up "
        "8% to 10% of the gross power generated. For example, on a day when gross generation at the generators "
        "sums up to 345 MKWh, only about 315 MKWh is successfully dispatched into the national grid. The rest is "
        "expended at the source as auxiliary station loss."
    )
    
    # SECTION 5: The Gas-Power Connection
    h5 = doc.add_heading(level=1)
    run_h5 = h5.add_run("5. The Gas-Power Connection")
    run_h5.font.name = "Georgia"
    run_h5.font.size = Pt(16)
    run_h5.font.bold = True
    run_h5.font.color.rgb = COLOR_PRIMARY
    
    p5 = doc.add_paragraph()
    p5.add_run(
        "Natural gas is Bangladesh's primary fuel source. When domestic gas fields deplete or global LNG prices "
        "skyrocket, Petrobangla is forced to cut the supply of gas to power plants. The Grid Explorer shows this "
        "cause-and-effect relationship visually:\n\n"
    )
    
    bp_gas1 = doc.add_paragraph(style='List Bullet')
    bp_gas1.add_run("Normal Operations: ").bold = True
    bp_gas1.add_run(
        "When gas supply is healthy (above 2,800 MMCFD), power plants receive their requested fuel, allowing "
        "them to generate electricity at full capacity. Load shedding is minimal, and electricity costs remain low."
    )
    
    bp_gas2 = doc.add_paragraph(style='List Bullet')
    bp_gas2.add_run("Deficit Scenarios: ").bold = True
    bp_gas2.add_run(
        "If gas supply drops to 2,200 MMCFD (a 600 MMCFD shortfall), gas-fired power plants are starved of fuel. "
        "To keep the lights on, the grid operator must switch on expensive diesel or furnace oil plants (driving up "
        "the bulk system cost in Crore Tk.) or cut power to regions entirely, causing several hundred Megawatts (MW) "
        "of load shedding."
    )
    
    # SECTION 6: Understanding the Historical Backlog
    h6 = doc.add_heading(level=1)
    run_h6 = h6.add_run("6. Understanding the Historical Backlog & UI Colors")
    run_h6.font.name = "Georgia"
    run_h6.font.size = Pt(16)
    run_h6.font.bold = True
    run_h6.font.color.rgb = COLOR_PRIMARY
    
    p6 = doc.add_paragraph()
    p6.add_run(
        "When exploring historical records, the Grid Explorer uses a color-shifting backlog badge to help you "
        "instantly identify the relative age of the archive date you are viewing. This serves as a psychological "
        "timeline indicator:"
    )
    
    # Table for backlog colors
    table_colors = doc.add_table(rows=1, cols=3)
    table_colors.alignment = WD_TABLE_ALIGNMENT.CENTER
    tc_hdr = table_colors.rows[0].cells
    tc_hdr[0].text = 'Time Difference'
    tc_hdr[1].text = 'Color Code'
    tc_hdr[2].text = 'Visual Meaning'
    
    for cell in tc_hdr:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = COLOR_PRIMARY
        
    colors_data = [
        ("0 - 90 Days", "Green", "Recent operational data; reflective of the current season and recent fuel prices."),
        ("90 - 365 Days", "Yellow", "Short-term archive; useful for comparing current metrics with the previous quarter."),
        ("1 - 3 Years", "Orange", "Medium-term historical archive; tracks changes across policy cycles."),
        ("3 - 7 Years", "Red", "Deep archive; shows grid performance before recent infrastructure upgrades."),
        ("Over 7 Years", "Rose / Crimson", "Deep historical records; shows the early baseline of Bangladesh's electricity grid expansion.")
    ]
    
    for duration, color, meaning in colors_data:
        row = table_colors.add_row().cells
        row[0].text = duration
        row[1].text = color
        row[2].text = meaning
        row[0].paragraphs[0].runs[0].font.bold = True
        
    doc.add_paragraph()
    
    # Footnote/Conclusion
    p_concl = doc.add_paragraph()
    r_concl = p_concl.add_run("Summary for the Reader:\n")
    r_concl.bold = True
    p_concl.add_run(
        "When you hand this sheet to anyone, they should understand: "
        "MW is the flow rate of power, MKWh is the daily accumulation of that power, MMCFD is the fuel flow "
        "enabling the generation, and the gaps between demand and supply dictate the loadshedding experienced in homes. "
        "With this foundation, the charts and numbers in the Power Grid Explorer will immediately become meaningful."
    )
    
    # Save document
    desktop_path = r"C:\Users\user\Desktop"
    output_filename = "Power_Grid_Explorer_Laymans_Guide.docx"
    full_path = os.path.join(desktop_path, output_filename)
    
    doc.save(full_path)
    print(f"Successfully saved handbook to {full_path}")

if __name__ == "__main__":
    generate_docx()
